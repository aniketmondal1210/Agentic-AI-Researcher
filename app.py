"""
Streamlit Frontend — AI Research Agent
Chat interface with multi-agent system, citation graphs, auth, and export options.
"""

import sys
import os
import json
import io
import re
from datetime import datetime

# Add the project root to path
sys.path.insert(0, os.path.dirname(__file__))

import streamlit as st
import streamlit.components.v1 as components
from agent.graph import INITIAL_PROMPT, graph, build_graph
from agent.multi_agent import build_multi_agent_graph, MULTI_AGENT_PROMPT
from tools.citation_graph import build_citation_graph_html
from pathlib import Path
import logging
import uuid
import yaml
from langchain_core.messages import AIMessage

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ─── Page Configuration ───
st.set_page_config(
    page_title="AI Research Agent",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── Custom CSS ───
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');

    .stApp { font-family: 'Inter', sans-serif; }

    /* ── Header ── */
    .main-header {
        background: #161b22;
        padding: 2rem;
        border-radius: 12px;
        margin-bottom: 1.5rem;
        text-align: center;
        border: 1px solid #30363d;
    }
    .main-header h1 {
        color: #e6edf3;
        font-size: 2rem;
        font-weight: 700;
        margin: 0;
    }
    .main-header p {
        color: #8b949e;
        font-size: 0.95rem;
        margin-top: 0.4rem;
        font-weight: 300;
    }

    /* ── Tool Call Box ── */
    .tool-call-box {
        background: #161b22;
        border-left: 3px solid #4fc3f7;
        border-radius: 6px;
        padding: 0.6rem 1rem;
        margin: 0.4rem 0;
        font-size: 0.85rem;
        color: #8b949e;
    }

    /* ── Agent Badges ── */
    .agent-badge {
        display: inline-block;
        padding: 0.2rem 0.6rem;
        border-radius: 12px;
        font-size: 0.75rem;
        font-weight: 600;
        margin-right: 0.4rem;
    }
    .agent-badge.search { background: rgba(79,195,247,0.12); color: #4fc3f7; }
    .agent-badge.reader { background: rgba(186,147,255,0.12); color: #ba93ff; }
    .agent-badge.writer { background: rgba(255,138,128,0.12); color: #ff8a80; }
    .agent-badge.supervisor { background: rgba(255,213,79,0.12); color: #ffd54f; }

    /* ── Legend Items ── */
    .legend-item {
        display: inline-block;
        margin: 0.2rem 0.4rem;
        padding: 0.25rem 0.6rem;
        border-radius: 8px;
        font-size: 0.8rem;
    }

    /* ── Info Cards ── */
    .info-card {
        background: #161b22;
        border: 1px solid #30363d;
        border-radius: 10px;
        padding: 0.8rem 1rem;
        margin: 0.4rem 0;
    }
    .info-card h4 {
        color: #e6edf3;
        margin: 0 0 0.2rem 0;
        font-size: 0.85rem;
        font-weight: 600;
    }
    .info-card p {
        color: #8b949e;
        margin: 0;
        font-size: 0.8rem;
    }

    /* ── Tabs ── */
    .stTabs [data-baseweb="tab-list"] {
        gap: 2px;
        background: #161b22;
        border-radius: 10px;
        padding: 3px;
        border: 1px solid #30363d;
    }
    .stTabs [data-baseweb="tab"] {
        border-radius: 8px;
        padding: 0.5rem 1rem;
        font-weight: 500;
        color: #8b949e;
    }
    .stTabs [aria-selected="true"] {
        background: #21262d !important;
        color: #e6edf3 !important;
        border: 1px solid #30363d !important;
    }
    .stTabs [data-baseweb="tab-highlight"] { display: none; }

    /* ── Buttons ── */
    .stButton > button {
        border-radius: 8px;
        font-weight: 500;
        border: 1px solid #30363d;
        transition: all 0.2s ease;
    }
    .stButton > button:hover { border-color: #4fc3f7; }

    /* ── Chat Messages ── */
    .stChatMessage { border-radius: 10px !important; }

    /* ── Metrics ── */
    [data-testid="stMetric"] {
        background: #161b22;
        border: 1px solid #30363d;
        border-radius: 10px;
        padding: 0.8rem;
    }

    /* ── Scrollbar ── */
    ::-webkit-scrollbar { width: 6px; }
    ::-webkit-scrollbar-track { background: transparent; }
    ::-webkit-scrollbar-thumb { background: #30363d; border-radius: 3px; }
    ::-webkit-scrollbar-thumb:hover { background: #484f58; }

    /* ── Hide defaults ── */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
</style>
""", unsafe_allow_html=True)


# ─── Authentication ───
def check_auth():
    """Handle user authentication."""
    try:
        import streamlit_authenticator as stauth

        # Load auth config
        config_path = str(Path(__file__).parent / "auth_config.yaml")

        authenticator = stauth.Authenticate(config_path)

        authenticator.login(location="main")

        if st.session_state.get("authentication_status"):
            authenticator.logout("🚪 Logout", location="sidebar")
            st.sidebar.markdown(f"**Welcome, {st.session_state.get('name', 'User')}!**")
            return True
        elif st.session_state.get("authentication_status") is False:
            st.error("❌ Invalid username or password")
            return False
        else:
            st.info("🔒 Please log in to use the AI Research Agent")
            st.markdown("**Default accounts:** `admin` / `admin123` or `researcher` / `research123`")
            return False

    except ImportError:
        # streamlit-authenticator not installed, skip auth
        return True
    except Exception as e:
        logger.warning(f"Auth error (skipping): {e}")
        return True


# Run auth check
if not check_auth():
    st.stop()


# ─── Helper Functions ───
def extract_text(content):
    """Extract clean text from AIMessage content."""
    if isinstance(content, str):
        return content
    elif isinstance(content, list):
        parts = []
        for block in content:
            if isinstance(block, dict) and block.get("type") == "text":
                parts.append(block["text"])
            elif isinstance(block, str):
                parts.append(block)
        return "\n".join(parts)
    return str(content)


def export_chat_markdown():
    """Export chat history as Markdown."""
    lines = [f"# AI Research Agent — Chat Export\n"]
    lines.append(f"**Date:** {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    lines.append(f"**Session:** {st.session_state.thread_id[:8]}\n\n---\n")
    for msg in st.session_state.chat_history:
        role = "🧑 User" if msg["role"] == "user" else "🤖 Assistant"
        lines.append(f"### {role}\n\n{msg['content']}\n\n---\n")
    return "\n".join(lines)


def export_chat_docx():
    """Export chat history as .docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading("AI Research Agent — Chat Export", level=0)
    doc.add_paragraph(
        f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')} | "
        f"Session: {st.session_state.thread_id[:8]}"
    )
    for msg in st.session_state.chat_history:
        heading = doc.add_heading(
            "User" if msg["role"] == "user" else "Assistant", level=2
        )
        run = heading.runs[0]
        run.font.color.rgb = RGBColor(50, 50, 150) if msg["role"] == "user" else RGBColor(0, 128, 80)
        para = doc.add_paragraph(msg["content"])
        for run in para.runs:
            run.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def extract_paper_titles_from_chat() -> list[str]:
    """Extract paper titles mentioned in the chat history."""
    titles = []
    for msg in st.session_state.chat_history:
        if msg["role"] == "assistant":
            content = msg["content"]
            # Match patterns like "**1. Title Here**" or "**Title:** Something"
            patterns = [
                r'\*\*\d+\.\s*(.+?)\*\*',
                r'\*\*Title:\*\*\s*(.+?)(?:\n|$)',
                r'\"(.{20,80})\"',
            ]
            for pattern in patterns:
                matches = re.findall(pattern, content)
                titles.extend(matches)
    # Deduplicate and limit
    seen = set()
    unique = []
    for t in titles:
        t_clean = t.strip().rstrip(".")
        if t_clean not in seen and len(t_clean) > 10:
            seen.add(t_clean)
            unique.append(t_clean)
    return unique[:10]


# ─── Session State ───
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "pdf_path" not in st.session_state:
    st.session_state.pdf_path = None
if "thread_id" not in st.session_state:
    st.session_state.thread_id = str(uuid.uuid4())
if "current_provider" not in st.session_state:
    st.session_state.current_provider = "gemini"
if "agent_mode" not in st.session_state:
    st.session_state.agent_mode = "single"
if "current_graph" not in st.session_state:
    st.session_state.current_graph = graph


# ─── Sidebar ───
with st.sidebar:
    st.markdown("## 🧠 Research Agent")
    st.markdown("---")

    # ── Agent Mode ──
    st.markdown("### ⚙️ Agent Mode")
    agent_mode = st.radio(
        "Mode",
        options=["single", "multi"],
        format_func=lambda x: "🤖 Single Agent (ReAct)" if x == "single" else "👥 Multi-Agent (Supervisor)",
        index=0 if st.session_state.agent_mode == "single" else 1,
        label_visibility="collapsed",
    )

    # ── LLM Provider ──
    st.markdown("### 🤖 LLM Provider")
    provider_options = {
        "gemini": "🔷 Google Gemini",
        "groq": "🟢 Groq (Llama)",
        "openai": "🟠 OpenAI (GPT)",
    }
    selected_provider = st.selectbox(
        "Choose LLM",
        options=list(provider_options.keys()),
        format_func=lambda x: provider_options[x],
        index=list(provider_options.keys()).index(st.session_state.current_provider),
        label_visibility="collapsed",
    )
    default_models = {
        "gemini": "gemini-2.0-flash",
        "groq": "llama-3.3-70b-versatile",
        "openai": "gpt-4o-mini",
    }
    model_name = st.text_input("Model name", value=default_models[selected_provider])

    # Rebuild graph if settings changed
    if (selected_provider != st.session_state.current_provider or
            agent_mode != st.session_state.agent_mode):
        try:
            if agent_mode == "multi":
                st.session_state.current_graph = build_multi_agent_graph(selected_provider, model_name)
            else:
                st.session_state.current_graph = build_graph(selected_provider, model_name)
            st.session_state.current_provider = selected_provider
            st.session_state.agent_mode = agent_mode
            st.success(f"Switched to {provider_options[selected_provider]} ({agent_mode} mode)")
        except Exception as e:
            st.error(f"Failed: {str(e)}")

    st.markdown("---")

    if st.button("🔄 New Session", use_container_width=True):
        st.session_state.chat_history = []
        st.session_state.pdf_path = None
        st.session_state.thread_id = str(uuid.uuid4())
        try:
            if st.session_state.agent_mode == "multi":
                st.session_state.current_graph = build_multi_agent_graph(
                    st.session_state.current_provider, model_name
                )
            else:
                st.session_state.current_graph = build_graph(
                    st.session_state.current_provider, model_name
                )
        except Exception:
            pass
        st.rerun()

    st.markdown(f"**Session:** `{st.session_state.thread_id[:8]}...`")
    st.markdown("---")

    # ── Upload Research Papers ──
    st.markdown("### 📎 Upload Papers")
    st.caption("Upload up to 6 PDFs — the agent can reference them")

    if "uploaded_papers" not in st.session_state:
        st.session_state.uploaded_papers = []

    uploaded_files = st.file_uploader(
        "Drop PDFs here",
        type=["pdf"],
        accept_multiple_files=True,
        key="pdf_uploader",
        label_visibility="collapsed",
    )

    if uploaded_files:
        # Limit to 6
        if len(uploaded_files) > 6:
            st.warning("⚠️ Maximum 6 papers allowed. Only the first 6 will be processed.")
            uploaded_files = uploaded_files[:6]

        new_files = [
            f for f in uploaded_files
            if f.name not in st.session_state.uploaded_papers
        ]

        if new_files:
            for uploaded_file in new_files:
                with st.spinner(f"📄 Indexing: {uploaded_file.name}..."):
                    try:
                        import PyPDF2

                        pdf_reader = PyPDF2.PdfReader(uploaded_file)
                        text = ""
                        for page in pdf_reader.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"

                        text = text.strip()
                        if not text:
                            st.warning(f"⚠️ Could not extract text from {uploaded_file.name}")
                            continue

                        # Truncate very long papers
                        if len(text) > 80000:
                            text = text[:80000]

                        # Store in RAG
                        from tools.rag_store import store_paper_in_rag
                        paper_title = uploaded_file.name.replace(".pdf", "").replace("_", " ").replace("-", " ")
                        result = store_paper_in_rag.invoke({
                            "paper_text": text,
                            "paper_title": paper_title,
                            "paper_url": f"uploaded:{uploaded_file.name}",
                        })

                        st.session_state.uploaded_papers.append(uploaded_file.name)
                        logger.info(f"Indexed uploaded paper: {uploaded_file.name} ({len(text)} chars)")

                    except Exception as e:
                        st.error(f"❌ Error processing {uploaded_file.name}: {str(e)}")

            if new_files:
                st.success(f"✅ {len(new_files)} paper(s) added to knowledge base!")
                st.rerun()

    # Show indexed uploads
    if st.session_state.uploaded_papers:
        for name in st.session_state.uploaded_papers:
            st.markdown(
                f"<div class='info-card'><h4>📄 {name}</h4><p>Indexed in knowledge base</p></div>",
                unsafe_allow_html=True,
            )
        if st.button("🗑️ Clear Uploads", use_container_width=True):
            st.session_state.uploaded_papers = []
            st.rerun()

    st.markdown("---")

    # ── Tools Info ──
    st.markdown("### 🛠️ Agent Tools")
    if agent_mode == "multi":
        st.markdown("""
<div class="info-card"><h4>🔍 Search Agent</h4><p>arXiv · Semantic Scholar · Web Search</p></div>
<div class="info-card"><h4>📖 Reader Agent</h4><p>PDF Reader · Summarizer · RAG Store</p></div>
<div class="info-card"><h4>✍️ Writer Agent</h4><p>LaTeX PDF · Quality Scorer · Lit Tables</p></div>
<div class="info-card"><h4>🧑‍💼 Supervisor</h4><p>Routes requests to the right specialist</p></div>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
<div class="info-card"><h4>🔍 Search</h4><p>arXiv · Semantic Scholar · Web</p></div>
<div class="info-card"><h4>📖 Read & Analyze</h4><p>PDF Reader · Summarizer · RAG Store</p></div>
<div class="info-card"><h4>✍️ Write & Export</h4><p>LaTeX PDF · Quality Scorer · Lit Tables</p></div>
        """, unsafe_allow_html=True)

    # ── Export ──
    if st.session_state.chat_history:
        st.markdown("---")
        st.markdown("### 📤 Export Chat")

        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                "📝 Markdown", export_chat_markdown(),
                f"chat_{st.session_state.thread_id[:8]}.md",
                "text/markdown", use_container_width=True,
            )
        with col2:
            st.download_button(
                "📋 JSON",
                json.dumps({
                    "session_id": st.session_state.thread_id,
                    "timestamp": datetime.now().isoformat(),
                    "provider": st.session_state.current_provider,
                    "mode": st.session_state.agent_mode,
                    "messages": st.session_state.chat_history,
                }, indent=2),
                f"chat_{st.session_state.thread_id[:8]}.json",
                "application/json", use_container_width=True,
            )
        try:
            st.download_button(
                "📄 Word (.docx)", export_chat_docx(),
                f"chat_{st.session_state.thread_id[:8]}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except ImportError:
            st.caption("Install `python-docx` for Word export")
        except Exception:
            pass

    # ── PDFs ──
    output_dir = Path("output")
    if output_dir.exists():
        pdf_files = list(output_dir.glob("*.pdf"))
        if pdf_files:
            st.markdown("---")
            st.markdown("### 📥 Generated Papers")
            for pdf_file in sorted(pdf_files, reverse=True)[:5]:
                with open(pdf_file, "rb") as f:
                    st.download_button(
                        f"⬇️ {pdf_file.name}", f.read(),
                        pdf_file.name, "application/pdf",
                        use_container_width=True,
                    )


# ─── Main Content ───
st.markdown("""
<div class="main-header">
    <h1>🔬 AI Research Agent</h1>
    <p>Search arXiv & Semantic Scholar · Read & Summarize Papers · RAG Knowledge Base · Write Publication-Ready LaTeX</p>
</div>
""", unsafe_allow_html=True)

# ─── Tabs ───
tab_chat, tab_citations, tab_knowledge, tab_lit_review = st.tabs(
    ["💬 Chat", "🔗 Citation Graph", "🧠 Knowledge Base", "📊 Literature Review"]
)

# ═══════════════════════════════════════════
#  TAB 1: Chat (with Human-in-the-Loop)
# ═══════════════════════════════════════════
with tab_chat:
    # Show agent mode badge
    if st.session_state.agent_mode == "multi":
        st.markdown(
            "<span class='agent-badge supervisor'>👥 Multi-Agent Mode</span>"
            "<span class='agent-badge search'>🔍 Search</span>"
            "<span class='agent-badge reader'>📖 Reader</span>"
            "<span class='agent-badge writer'>✍️ Writer</span>",
            unsafe_allow_html=True,
        )

    # ── Human-in-the-Loop Controls ──
    if "hitl_enabled" not in st.session_state:
        st.session_state.hitl_enabled = False
    if "hitl_pending" not in st.session_state:
        st.session_state.hitl_pending = None

    with st.expander("🔧 Human-in-the-Loop Settings", expanded=False):
        st.session_state.hitl_enabled = st.toggle(
            "Enable approval checkpoints",
            value=st.session_state.hitl_enabled,
            help="When enabled, the agent will pause for your approval before writing or generating PDFs.",
        )
        if st.session_state.hitl_enabled:
            st.info("✅ The agent will ask for your approval before: writing papers, generating PDFs, and scoring quality.")

    # Show pending approval if any
    if st.session_state.hitl_pending:
        st.warning("⏸️ **Approval Required**")
        st.markdown(f"The agent wants to: **{st.session_state.hitl_pending['action']}**")
        if st.session_state.hitl_pending.get("details"):
            with st.expander("View details"):
                st.markdown(st.session_state.hitl_pending["details"])
        col1, col2, col3 = st.columns(3)
        with col1:
            if st.button("✅ Approve", use_container_width=True, type="primary"):
                # Add approval to chat and continue
                st.session_state.chat_history.append(
                    {"role": "user", "content": f"Approved: {st.session_state.hitl_pending['action']}. Go ahead."}
                )
                st.session_state.hitl_pending = None
                st.rerun()
        with col2:
            if st.button("✏️ Modify", use_container_width=True):
                st.session_state.hitl_pending = None
                st.rerun()
        with col3:
            if st.button("❌ Reject", use_container_width=True):
                st.session_state.chat_history.append(
                    {"role": "user", "content": f"Rejected: {st.session_state.hitl_pending['action']}. Please try a different approach."}
                )
                st.session_state.hitl_pending = None
                st.rerun()

    # Display chat history
    for msg in st.session_state.chat_history:
        if msg["role"] in ("user", "assistant"):
            st.chat_message(msg["role"]).write(msg["content"])

    # Chat input
    user_input = st.chat_input("What research topic would you like to explore?")

    if user_input:
        logger.info(f"User input: {user_input}")
        st.session_state.chat_history.append({"role": "user", "content": user_input})
        st.chat_message("user").write(user_input)

        # Choose the right prompt
        system_prompt = (
            MULTI_AGENT_PROMPT if st.session_state.agent_mode == "multi"
            else INITIAL_PROMPT
        )

        # Add HITL instructions if enabled
        if st.session_state.hitl_enabled:
            system_prompt += (
                "\n\nIMPORTANT: Before writing a paper or generating a PDF, "
                "first describe your plan and ASK THE USER FOR APPROVAL. "
                "Present an outline of what you plan to write and wait for the user to confirm."
            )

        # Add uploaded papers context
        if st.session_state.get("uploaded_papers"):
            papers_list = ", ".join(st.session_state.uploaded_papers)
            system_prompt += (
                f"\n\nThe user has uploaded these research papers to the knowledge base: {papers_list}. "
                "You can use query_rag_store to search through their content and reference them. "
                "When the user asks about these papers, always query the knowledge base first."
            )

        chat_input_data = {
            "messages": [
                {"role": "system", "content": system_prompt}
            ] + st.session_state.chat_history
        }

        if st.session_state.agent_mode == "multi":
            chat_input_data["current_agent"] = ""
            chat_input_data["task_complete"] = False

        config = {"configurable": {"thread_id": st.session_state.thread_id}}

        full_response = ""
        with st.spinner("🤖 Agent is working..."):
            try:
                for s in st.session_state.current_graph.stream(
                    chat_input_data, config, stream_mode="values"
                ):
                    message = s["messages"][-1]

                    if getattr(message, "tool_calls", None):
                        for tc in message.tool_calls:
                            logger.info(f"Tool call: {tc['name']}")

                            # HITL: Pause for approval on critical tools
                            if (st.session_state.hitl_enabled and
                                    tc["name"] in ("render_latex_pdf", "score_paper_quality")):
                                st.session_state.hitl_pending = {
                                    "action": f"Use {tc['name']}",
                                    "details": f"Tool: {tc['name']}\nArgs: {json.dumps(tc.get('args', {}), indent=2)[:500]}",
                                }

                            st.markdown(
                                f"<div class='tool-call-box'>🔧 Using tool: <b>{tc['name']}</b></div>",
                                unsafe_allow_html=True,
                            )

                    if "current_agent" in s and s["current_agent"]:
                        agent_name = s["current_agent"]
                        badge_class = {
                            "search_agent": "search",
                            "reader_agent": "reader",
                            "writer_agent": "writer",
                        }.get(agent_name, "supervisor")
                        st.markdown(
                            f"<div class='tool-call-box'>"
                            f"<span class='agent-badge {badge_class}'>"
                            f"Routing to: {agent_name}</span></div>",
                            unsafe_allow_html=True,
                        )

                    if isinstance(message, AIMessage) and message.content:
                        text_content = extract_text(message.content)
                        if text_content.strip():
                            full_response = text_content

                if full_response:
                    st.chat_message("assistant").write(full_response)
                    st.session_state.chat_history.append(
                        {"role": "assistant", "content": full_response}
                    )

                # Check for new PDFs
                output_dir = Path("output")
                if output_dir.exists():
                    pdf_files = sorted(output_dir.glob("*.pdf"), reverse=True)
                    if pdf_files:
                        latest = pdf_files[0]
                        if st.session_state.pdf_path != str(latest):
                            st.session_state.pdf_path = str(latest)
                            st.success(f"📄 New PDF: {latest.name}")
                            with open(latest, "rb") as f:
                                st.download_button(
                                    f"⬇️ Download {latest.name}",
                                    f.read(), latest.name,
                                    "application/pdf",
                                    use_container_width=True,
                                )

            except Exception as e:
                st.error(f"❌ An error occurred: {str(e)}")
                logger.error(f"Agent error: {str(e)}", exc_info=True)


# ═══════════════════════════════════════════
#  TAB 2: Citation Graph
# ═══════════════════════════════════════════
with tab_citations:
    st.markdown("### 🔗 Citation Network Graph")
    st.markdown(
        "Visualize how papers cite each other. "
        "Enter paper titles below or auto-extract from the chat."
    )

    auto_titles = extract_paper_titles_from_chat()

    col1, col2 = st.columns([3, 1])
    with col1:
        paper_input = st.text_area(
            "Paper titles (one per line)",
            value="\n".join(auto_titles[:5]) if auto_titles else "",
            height=120,
            placeholder="Enter paper titles, one per line...",
        )
    with col2:
        st.markdown("")
        st.markdown("")
        build_btn = st.button("🔍 Build Graph", use_container_width=True, type="primary")

    st.markdown(
        "<div>"
        "<span class='legend-item' style='background:rgba(231,76,60,0.2);color:#e74c3c;'>● Main Papers</span>"
        "<span class='legend-item' style='background:rgba(0,184,148,0.2);color:#00b894;'>● Citing Papers</span>"
        "<span class='legend-item' style='background:rgba(108,92,231,0.2);color:#6c5ce7;'>● References</span>"
        "</div>",
        unsafe_allow_html=True,
    )

    if build_btn and paper_input.strip():
        titles = [t.strip() for t in paper_input.strip().split("\n") if t.strip()]
        if titles:
            with st.spinner("⏳ Fetching citation data from Semantic Scholar..."):
                try:
                    html = build_citation_graph_html(titles, height="550px")
                    components.html(html, height=600, scrolling=True)
                except Exception as e:
                    st.error(f"Error building graph: {str(e)}")
        else:
            st.warning("Please enter at least one paper title.")
    elif build_btn:
        st.warning("Please enter paper titles first.")


# ═══════════════════════════════════════════
#  TAB 3: Knowledge Base (RAG)
# ═══════════════════════════════════════════
with tab_knowledge:
    st.markdown("### 🧠 Knowledge Base (RAG Vector Store)")
    st.markdown("Papers read by the agent are stored here for semantic search.")

    try:
        from tools.rag_store import get_rag_stats

        stats = get_rag_stats()

        # Stats cards
        col1, col2 = st.columns(2)
        with col1:
            st.metric("📄 Papers Indexed", stats["total_papers"])
        with col2:
            st.metric("🧩 Total Chunks", stats["total_chunks"])

        if stats["paper_titles"]:
            st.markdown("#### 📚 Indexed Papers")
            for title in stats["paper_titles"]:
                st.markdown(f"- {title}")

        # Query interface
        st.markdown("---")
        st.markdown("#### 🔍 Search Knowledge Base")
        rag_query = st.text_input(
            "Search query",
            placeholder="Search for specific concepts, methods, or findings...",
            key="rag_query_input",
        )

        if rag_query:
            from tools.rag_store import query_rag_store

            with st.spinner("Searching knowledge base..."):
                results = query_rag_store.invoke({"query": rag_query, "n_results": 5})
                st.markdown(results)

    except ImportError:
        st.warning("ChromaDB not installed. Run: `pip install chromadb`")
    except Exception as e:
        st.info(f"Knowledge base is empty. Start chatting with the agent and ask it to read papers!")


# ═══════════════════════════════════════════
#  TAB 4: Literature Review Table
# ═══════════════════════════════════════════
with tab_lit_review:
    st.markdown("### 📊 Automated Literature Review Table")
    st.markdown(
        "Enter information about multiple papers to generate a structured comparison table. "
        "You can also ask the agent in the chat tab to generate one."
    )

    lit_input = st.text_area(
        "Paper information",
        height=200,
        placeholder=(
            "Paste paper details here. For each paper include:\n"
            "- Title\n- Authors\n- Year\n- Methodology\n"
            "- Key findings\n- Dataset used\n\n"
            "Or simply paste the text from multiple paper abstracts."
        ),
        key="lit_review_input",
    )

    if st.button("📊 Generate Table", use_container_width=False, type="primary"):
        if lit_input.strip():
            with st.spinner("🤖 Generating literature review table..."):
                try:
                    from tools.literature_table import generate_literature_table

                    prompt = generate_literature_table.invoke({"papers_info": lit_input})

                    # Use the current LLM to generate the table (without tool binding)
                    from agent.graph import get_model
                    llm = get_model(st.session_state.current_provider, bind=False)
                    response = llm.invoke(prompt)

                    table_content = extract_text(response.content)
                    st.markdown(table_content)

                    # Download as markdown
                    st.download_button(
                        "📥 Download Table (Markdown)",
                        table_content,
                        "literature_review_table.md",
                        "text/markdown",
                    )
                except Exception as e:
                    st.error(f"Error generating table: {str(e)}")
        else:
            st.warning("Please enter paper information first.")

